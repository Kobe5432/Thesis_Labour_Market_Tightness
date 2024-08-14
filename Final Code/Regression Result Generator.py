import pandas as pd
from docx import Document

# Load the Excel file
file_path = r'regression_summary.xlsx'
df = pd.read_excel(file_path)

# Define the desired order of variables
desired_order = ['lagged_inflation_expectations','lagged_ma_productivity_inflation', 'lagged_ma_job_vacancy_rate', 'const']

# Add significance stars based on p-values
def add_stars(row):
    if row['p-value'] < 0.01:
        return '***'
    elif row['p-value'] < 0.05:
        return '**'
    elif row['p-value'] < 0.1:
        return '*'
    else:
        return ''

# Format numerical values to scientific notation with 3 significant digits
def format_scientific(value, is_r_squared=False, is_coefficient=False, is_observations=False):
    if pd.isna(value):
        return ''
    try:
        float_value = float(value)
        if is_r_squared:
            return f"{float_value:.2f}"  # Format R-squared values to three decimal places
        elif is_coefficient:
            return f"{float_value:.3f}" 
        elif is_observations:
            return f"{float_value:.0f}" 
        else:
            return f"{float_value:.3f}"
    
    except (ValueError, TypeError):
        return str(value)

# Create columns for combined model data
df['ModelData'] = df.apply(lambda row: f"{format_scientific(row['Coefficient'], is_coefficient=True)}{add_stars(row)}\n({format_scientific(row['Standard Error'])})", axis=1)

# Extract R-squared values
r_squared = df[['Section', 'Model', 'Adjusted R-squared']].drop_duplicates().pivot(index='Section', columns='Model', values='Adjusted R-squared')

# Extract R-squared values
observations = df[['Section', 'Model', 'Number of Observations']].drop_duplicates().pivot(index='Section', columns='Model', values='Number of Observations')

# Pivot the table to get the desired format
pivot_table = df.pivot(index=['Section', 'Variable'], columns='Model', values='ModelData')
pivot_table.columns = [f"{col}" for col in pivot_table.columns]

# Create a Word document
doc = Document()

# Add a title
doc.add_heading('Regression Results', level=1)

# Iterate over unique sections and create tables
for section in pivot_table.index.get_level_values('Section').unique():
    # Filter pivot_table for the current section and sort variables in desired order
    section_data = pivot_table.loc[section].reindex(desired_order)
    
    # Create a table in the Word document for the current section
    doc.add_heading(f'Economic Activity: {section}', level=2)
    
    table = doc.add_table(rows=len(section_data) + 2, cols=len(section_data.columns) + 1)  # Adjusted rows to add extra row for observations
    table.style = 'Table Grid'
    
    # Add the header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Variable'
    for i, col in enumerate(section_data.columns):
        hdr_cells[i + 1].text = col
    
    # Add the data rows
    for idx, row in section_data.iterrows():
        variable = idx  # Use the second element of the index tuple as the variable name
        row_cells = table.add_row().cells
        row_cells[0].text = variable  # First cell in each row is the variable
        for i, value in enumerate(row):
            if pd.isna(value):
                row_cells[i + 1].text = ''  # Replace NaN with empty cell
            else:
                row_cells[i + 1].text = format_scientific(value)
    
    # Delete the first 5 empty rows if they exist
    while len(table.rows) > 1 and table.rows[1].cells[0].text == '':
        table._element.remove(table.rows[1]._element)
    
    # Add extra row for R-squared values
    observations_row = table.add_row().cells
    observations_row[0].text = 'Number of Observations'

    # Debugging: print if the section exists in R-squared DataFrame
    if section in observations.index:
        for i, model in enumerate(observations.columns):
            observations_row[i + 1].text = format_scientific(observations.loc[section, model], is_observations=True)
    else:
        for i in range(1, len(section_data.columns) + 1):
            observations_row[i].text = ''

    # Add extra row for R-squared values
    r_squared_row = table.add_row().cells
    r_squared_row[0].text = 'R-squared'
    
    # Debugging: print if the section exists in R-squared DataFrame
    if section in r_squared.index:
        print(f"R-squared values found for section: {section}")
        for i, model in enumerate(r_squared.columns):
            r_squared_row[i + 1].text = format_scientific(r_squared.loc[section, model], is_r_squared=True)
    else:
        print(f"No R-squared values found for section: {section}")
        for i in range(1, len(section_data.columns) + 1):
            r_squared_row[i].text = ''

# Save the document
output_path_docx  = r'regression_results_tables.docx'
doc.save(output_path_docx )

print(f"Tables successfully saved to {output_path_docx }")